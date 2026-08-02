from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph


def extract_text(file_path: Path) -> str:
    """
    Reads a .docx file's body and returns its text, in reading order.

    `Document.iter_inner_content()` walks the body's paragraphs *and*
    tables in the order they actually appear (unlike `.paragraphs` /
    `.tables`, which are two separate flat lists and lose where a
    table sat relative to the paragraphs around it) — that's what
    "preserving reading order" means here.

    Mirrors pdf_service.extract_text's shape (one function, a `Path`
    in, a `str` out) and joins blocks with blank lines between them
    for the same reason pdf_service joins pages that way: it gives the
    LLM a visible break between distinct chunks of the document
    instead of one run-on wall of text.

    Note: like pdf_service, this only reads text already in the
    document — text inside images (e.g. a scanned page pasted in as a
    picture) has no text layer here either, and is out of scope for
    the same OCR reason noted there.
    """
    document = DocxDocument(str(file_path))

    blocks: list[str] = []
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if text:
                blocks.append(text)
        elif isinstance(item, Table):
            table_text = _table_to_text(item)
            if table_text:
                blocks.append(table_text)

    return "\n\n".join(blocks).strip()


def _table_to_text(table: Table) -> str:
    """
    Renders a table as one line per row, cells separated by " | ", so
    a row's cells stay readable (and stay grouped together) as plain
    text instead of losing all structure. Rows that are entirely blank
    (e.g. spacer rows) are dropped.
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)
