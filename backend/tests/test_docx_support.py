"""
DOCX support (V2.2 Milestone 1).

Mirrors the existing PDF test files' shape and conventions
(test_documents.py, test_rag.py, test_multi_document_chat.py, ...) —
these tests exist to prove the *same* pipeline (upload -> extraction
-> chunking -> embeddings -> RAG -> summary/flashcards/quiz/mindmap/
chat) behaves identically for a .docx-sourced document as it already
does for a .pdf-sourced one, not to reimplement coverage of the
pipeline itself.
"""

import io
import json

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
from fastapi.testclient import TestClient
from reportlab.pdfgen import canvas

from app.main import app
from app.services import docx_service, document_extraction_service
from app.services.ai.base_provider import AIProvider, AIProviderError
from app.services.ai.embedding_provider import EmbeddingProvider
from app.services.ai.embedding_provider_factory import get_embedding_provider
from app.services.ai.provider_factory import get_ai_provider

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

client = TestClient(app)


# ---------------------------------------------------------------------------
# Fakes (local to this file, same convention every other test file uses —
# see docs/architecture.md's "Testing conventions").
# ---------------------------------------------------------------------------


class FakeAIProvider(AIProvider):
    """Stands in for a real provider — instant, free, deterministic."""

    async def generate_text(self, prompt: str) -> str:
        return "This is a fake summary for testing."


class FakeStructuredAIProvider(AIProvider):
    """Returns whatever well-formed JSON string it's constructed with,
    for flashcards/quiz/mindmap, which all expect structured output."""

    def __init__(self, response: str) -> None:
        self._response = response

    async def generate_text(self, prompt: str) -> str:
        return self._response


class FakeEmbeddingProvider(EmbeddingProvider):
    """Deterministic embeddings, same technique as test_rag.py."""

    async def embed_document(self, text: str) -> list[float]:
        return self._vector_for(text)

    async def embed_query(self, text: str) -> list[float]:
        return self._vector_for(text)

    @staticmethod
    def _vector_for(text: str) -> list[float]:
        lowered = text.lower()
        cat_count = lowered.count("cat")
        dog_count = lowered.count("dog")
        if cat_count > dog_count:
            return [1.0, 0.0]
        if dog_count > cat_count:
            return [0.0, 1.0]
        return [0.5, 0.5]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_test_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """
    Builds a tiny real .docx in memory, the same "generate a real file
    rather than depend on a fixture on disk" approach _make_test_pdf
    (reportlab) uses for PDFs. `table_rows`, if given, is inserted
    between the first and second paragraph so tests can prove reading
    order is preserved across a table.
    """
    document = docx.Document()
    document.add_paragraph(paragraphs[0])
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row_values in enumerate(table_rows):
            for col_index, value in enumerate(row_values):
                table.cell(row_index, col_index).text = value
    for paragraph in paragraphs[1:]:
        document.add_paragraph(paragraph)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_test_pdf(text: str) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(100, 750, text)
    pdf.save()
    return buffer.getvalue()


def _upload_docx(filename: str, docx_bytes: bytes):
    return client.post(
        "/api/v1/documents/upload",
        files={"file": (filename, docx_bytes, DOCX_CONTENT_TYPE)},
    )


def _upload_ready_docx(filename: str = "test.docx", text: str = "Some content.") -> str:
    response = _upload_docx(filename, _make_test_docx([text]))
    return response.json()["id"]


def _upload_and_index_docx(filename: str, text: str) -> str:
    document_id = _upload_ready_docx(filename, text)
    client.post(f"/api/v1/documents/{document_id}/index")
    return document_id


# ---------------------------------------------------------------------------
# docx_service — unit-level extraction behavior
# ---------------------------------------------------------------------------


def test_docx_extract_text_reads_paragraphs(tmp_path):
    path = tmp_path / "notes.docx"
    path.write_bytes(_make_test_docx(["First paragraph.", "Second paragraph."]))

    text = docx_service.extract_text(path)

    assert "First paragraph." in text
    assert "Second paragraph." in text
    assert text.index("First paragraph.") < text.index("Second paragraph.")


def test_docx_extract_text_preserves_reading_order_across_a_table(tmp_path):
    """
    A table placed between two paragraphs must come back between them
    in the extracted text, not before/after both — this is the
    "preserve reading order" requirement, and the reason
    docx_service.py uses `iter_inner_content()` instead of the
    separate `.paragraphs` / `.tables` lists.
    """
    path = tmp_path / "with_table.docx"
    path.write_bytes(
        _make_test_docx(
            ["Intro paragraph.", "Outro paragraph."],
            table_rows=[["Header A", "Header B"], ["Cell A1", "Cell B1"]],
        )
    )

    text = docx_service.extract_text(path)

    intro_pos = text.index("Intro paragraph.")
    table_pos = text.index("Header A")
    outro_pos = text.index("Outro paragraph.")
    assert intro_pos < table_pos < outro_pos
    # Cells from the same row stay together on one line.
    assert "Header A | Header B" in text
    assert "Cell A1 | Cell B1" in text


def test_docx_extract_text_returns_empty_string_for_blank_document(tmp_path):
    path = tmp_path / "blank.docx"
    path.write_bytes(_make_test_docx([""]))

    assert docx_service.extract_text(path) == ""


def test_docx_extract_text_raises_for_a_corrupted_file(tmp_path):
    path = tmp_path / "corrupted.docx"
    path.write_bytes(b"this is not a real docx file")

    try:
        docx_service.extract_text(path)
        assert False, "expected extraction to raise for a corrupted file"
    except Exception:
        pass


# ---------------------------------------------------------------------------
# document_extraction_service — dispatch abstraction
# ---------------------------------------------------------------------------


def test_extraction_service_dispatches_docx_to_docx_service(tmp_path):
    path = tmp_path / "routed.docx"
    path.write_bytes(_make_test_docx(["Routed through the dispatcher."]))

    text = document_extraction_service.extract_text(path, ".docx")

    assert "Routed through the dispatcher." in text


def test_extraction_service_has_no_page_counter_for_docx(tmp_path):
    path = tmp_path / "no_pages.docx"
    path.write_bytes(_make_test_docx(["Some content."]))

    assert document_extraction_service.get_page_count(path, ".docx") is None


def test_extraction_service_extension_matching_is_case_insensitive(tmp_path):
    path = tmp_path / "upper.docx"
    path.write_bytes(_make_test_docx(["Case insensitive."]))

    text = document_extraction_service.extract_text(path, ".DOCX")

    assert "Case insensitive." in text


def test_extraction_service_raises_for_unsupported_extension(tmp_path):
    path = tmp_path / "sheet.xlsx"
    path.write_bytes(b"not really an xlsx")

    try:
        document_extraction_service.extract_text(path, ".xlsx")
        assert False, "expected UnsupportedFileTypeError"
    except document_extraction_service.UnsupportedFileTypeError:
        pass


# ---------------------------------------------------------------------------
# Upload endpoint
# ---------------------------------------------------------------------------


def test_upload_docx_extracts_text():
    response = _upload_docx("hello.docx", _make_test_docx(["Hello LearnFlow"]))

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "ready"
    assert "Hello LearnFlow" in body["text_preview"]
    assert body["character_count"] > 0


def test_upload_docx_includes_file_size_and_null_page_count():
    docx_bytes = _make_test_docx(["One doc, no page count."])

    response = _upload_docx("sized.docx", docx_bytes)

    assert response.status_code == 201
    body = response.json()
    assert body["file_size_bytes"] == len(docx_bytes)
    # DOCX has no reliable page tree the way PDF does — the field
    # stays present (like it does for every document) but null.
    assert body["page_count"] is None


def test_upload_docx_preserves_original_filename_and_extension():
    response = _upload_docx("Lecture Notes.docx", _make_test_docx(["Content."]))

    assert response.status_code == 201
    assert response.json()["original_filename"] == "Lecture Notes.docx"


def test_upload_rejects_empty_docx_file():
    response = _upload_docx("empty.docx", b"")

    assert response.status_code == 400


def test_upload_docx_with_corrupted_bytes_uploads_but_fails_processing():
    """
    A corrupted .docx isn't rejected at the HTTP boundary (the
    content-type is legitimately "this is a DOCX upload") — same
    behavior as a corrupted PDF today: the record is created so the
    user can see it happened, and status flips to "failed" rather than
    a 500 or a silently empty document.
    """
    response = _upload_docx("corrupted.docx", b"this is not a real docx file")

    assert response.status_code == 201
    assert response.json()["status"] == "failed"

    get_response = client.get(f"/api/v1/documents/{response.json()['id']}")
    assert get_response.json()["status"] == "failed"


def test_upload_rejects_unsupported_file_type():
    response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "sheet.xlsx",
                b"not supported",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )

    assert response.status_code == 400
    assert "PDF, DOCX, PPTX, PNG, JPG, and JPEG" in response.json()["detail"]


def test_get_docx_document_returns_it_after_upload():
    upload_response = _upload_docx("retrieve.docx", _make_test_docx(["Retrieve me"]))
    document_id = upload_response.json()["id"]

    get_response = client.get(f"/api/v1/documents/{document_id}")

    assert get_response.status_code == 200
    assert get_response.json()["id"] == document_id


def test_pdf_and_docx_documents_coexist_in_the_library():
    pdf_response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("mixed-lib.pdf", _make_test_pdf("A PDF."), "application/pdf")},
    )
    docx_response = _upload_docx("mixed-lib.docx", _make_test_docx(["A DOCX."]))

    list_response = client.get("/api/v1/documents", params={"search": "mixed-lib"})

    assert list_response.status_code == 200
    names = {doc["original_filename"] for doc in list_response.json()}
    assert {"mixed-lib.pdf", "mixed-lib.docx"} <= names
    assert pdf_response.json()["status"] == "ready"
    assert docx_response.json()["status"] == "ready"


def test_rename_and_delete_round_trip_for_an_uploaded_docx():
    document_id = _upload_ready_docx("Original.docx", "Some content.")

    rename_response = client.patch(
        f"/api/v1/documents/{document_id}", json={"original_filename": "Renamed"}
    )
    assert rename_response.status_code == 200
    assert rename_response.json()["original_filename"] == "Renamed.docx"

    delete_response = client.delete(f"/api/v1/documents/{document_id}")
    assert delete_response.status_code == 204

    get_response = client.get(f"/api/v1/documents/{document_id}")
    assert get_response.status_code == 404


# ---------------------------------------------------------------------------
# RAG (chunking + embeddings + retrieval) for a DOCX-sourced document
# ---------------------------------------------------------------------------


def test_index_and_search_a_docx_document():
    app.dependency_overrides[get_embedding_provider] = lambda: FakeEmbeddingProvider()
    try:
        text = ("The cat sat on the mat. " * 60) + ("The dog ran in the yard. " * 60)
        document_id = _upload_and_index_docx("cats_and_dogs.docx", text)

        index_response = client.get(f"/api/v1/documents/{document_id}")
        assert index_response.json()["status"] == "ready"

        search_response = client.post(
            f"/api/v1/documents/{document_id}/search", json={"query": "Tell me about cats"}
        )
        assert search_response.status_code == 200
        results = search_response.json()["results"]
        assert results
        assert "cat" in results[0]["content"].lower()
    finally:
        app.dependency_overrides.clear()


# ---------------------------------------------------------------------------
# Summary / Flashcards / Quiz / Mind Map for a DOCX-sourced document
# ---------------------------------------------------------------------------


def test_summary_works_for_a_docx_document():
    app.dependency_overrides[get_ai_provider] = lambda: FakeAIProvider()
    try:
        document_id = _upload_ready_docx("summary.docx", "Photosynthesis converts light to energy.")

        response = client.post(f"/api/v1/documents/{document_id}/summary")

        assert response.status_code == 201
        assert response.json()["content"] == "This is a fake summary for testing."
    finally:
        app.dependency_overrides.clear()


def test_flashcards_work_for_a_docx_document():
    fake_cards = json.dumps([{"question": "What is 2+2?", "answer": "4"}])
    app.dependency_overrides[get_ai_provider] = lambda: FakeStructuredAIProvider(fake_cards)
    try:
        document_id = _upload_ready_docx("flashcards.docx", "Basic arithmetic.")

        response = client.post(f"/api/v1/documents/{document_id}/flashcards")

        assert response.status_code == 201
        cards = response.json()
        assert len(cards) == 1
        assert cards[0]["question"] == "What is 2+2?"
    finally:
        app.dependency_overrides.clear()


def test_quiz_works_for_a_docx_document():
    fake_quiz = json.dumps(
        [
            {
                "question": "What is the capital of France?",
                "options": ["Paris", "Rome", "Berlin", "Madrid"],
                "correct_answer_index": 0,
            }
        ]
    )
    app.dependency_overrides[get_ai_provider] = lambda: FakeStructuredAIProvider(fake_quiz)
    try:
        document_id = _upload_ready_docx("quiz.docx", "European capitals.")

        response = client.post(f"/api/v1/documents/{document_id}/quiz")

        assert response.status_code == 201
        questions = response.json()
        assert len(questions) == 1
        assert questions[0]["correct_answer_index"] == 0
    finally:
        app.dependency_overrides.clear()


def test_mindmap_works_for_a_docx_document():
    fake_mindmap = json.dumps({"title": "Central Topic", "children": []})
    app.dependency_overrides[get_ai_provider] = lambda: FakeStructuredAIProvider(fake_mindmap)
    try:
        document_id = _upload_ready_docx("mindmap.docx", "A topic with no subtopics.")

        response = client.post(f"/api/v1/documents/{document_id}/mindmap")

        assert response.status_code == 201
        assert response.json()["structure"]["title"] == "Central Topic"
    finally:
        app.dependency_overrides.clear()


# ---------------------------------------------------------------------------
# Chat (single- and multi-document) for a DOCX-sourced document
# ---------------------------------------------------------------------------


class FakeChatProvider(AIProvider):
    async def generate_text(self, prompt: str) -> str:
        return "Cats are mentioned in the excerpts."


def test_single_document_chat_works_for_a_docx_document():
    app.dependency_overrides[get_embedding_provider] = lambda: FakeEmbeddingProvider()
    app.dependency_overrides[get_ai_provider] = lambda: FakeChatProvider()
    try:
        document_id = _upload_and_index_docx("chat.docx", "The cat sat on the mat. " * 60)

        response = client.post(
            f"/api/v1/documents/{document_id}/chat", json={"question": "Tell me about cats"}
        )

        assert response.status_code == 200
        body = response.json()
        assert body["grounded"] is True
        assert body["answer"] == "Cats are mentioned in the excerpts."
    finally:
        app.dependency_overrides.clear()


def test_multi_document_chat_works_across_a_pdf_and_a_docx_document():
    """
    Proves multi-document chat — the code path the architecture doc
    calls out as shared between single- and multi-document chat — is
    genuinely format-agnostic: one PDF-sourced document and one
    DOCX-sourced document, selected together, both contribute.
    """
    app.dependency_overrides[get_embedding_provider] = lambda: FakeEmbeddingProvider()
    app.dependency_overrides[get_ai_provider] = lambda: FakeChatProvider()
    try:
        pdf_upload = client.post(
            "/api/v1/documents/upload",
            files={
                "file": (
                    "cats.pdf",
                    _make_test_pdf("The cat sat on the mat. " * 30),
                    "application/pdf",
                )
            },
        )
        pdf_id = pdf_upload.json()["id"]
        client.post(f"/api/v1/documents/{pdf_id}/index")

        docx_id = _upload_and_index_docx("dogs.docx", "The dog ran in the yard. " * 30)

        response = client.post(
            "/api/v1/documents/chat",
            json={"document_ids": [pdf_id, docx_id], "question": "Compare cats and dogs"},
        )

        assert response.status_code == 200
        body = response.json()
        sources_by_document = {source["document_id"] for source in body["sources"]}
        assert pdf_id in sources_by_document
        assert docx_id in sources_by_document
    finally:
        app.dependency_overrides.clear()
